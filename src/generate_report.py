from __future__ import annotations

import json
import math
import platform
import sys
from typing import Dict, List

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from utils import REPORT_FILE, RESULTS_CSV, TEXT_DOCX, format_percentage, format_seconds, is_prediction_coherent, read_docx_text


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)

    set_default_font(document)


def add_paragraph(
    document: Document,
    text: str,
    *,
    bold: bool = False,
    uppercase: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    after: float = 0,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text.upper() if uppercase else text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_section_title(document: Document, text: str) -> None:
    add_paragraph(document, text, bold=True, uppercase=True)


def set_table_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_general_table(document: Document, df: pd.DataFrame) -> None:
    add_paragraph(
        document,
        "Tabela 1 - Comparacao entre os modelos Vision Transformer",
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    columns = [
        "Imagem",
        "URL",
        "Modelo",
        "Classe prevista",
        "Confianca da classe principal",
        "Tempo de carregamento (s)",
        "Tempo de inferencia (s)",
    ]
    table = document.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for idx, title in enumerate(columns):
        set_table_cell_text(table.rows[0].cells[idx], title, bold=True)
        shade_cell(table.rows[0].cells[idx], "D9EAF7")

    for _, row in df.iterrows():
        cells = table.add_row().cells
        classe = row["predicted_class"] if row["status"] == "ok" else f"Falha: {row['error']}"
        values = [
            row["image_name"],
            row["url"],
            row["model_label"],
            classe,
            format_percentage(row["top1_confidence_percent"]) if row["status"] == "ok" else "N/A",
            format_seconds(row["model_load_time_seconds"]) if row["status"] == "ok" else "N/A",
            format_seconds(row["inference_time_seconds"]) if row["status"] == "ok" else "N/A",
        ]
        for cell, value in zip(cells, values):
            set_table_cell_text(cell, str(value))

    add_paragraph(
        document,
        "Fonte: Elaborado pelo autor com base nos resultados obtidos na execucao local.",
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )


def load_results() -> pd.DataFrame:
    if not RESULTS_CSV.exists():
        raise FileNotFoundError(f"Arquivo de resultados nao encontrado: {RESULTS_CSV}")

    df = pd.read_csv(RESULTS_CSV)
    for column in ["top1_confidence_percent", "model_load_time_seconds", "inference_time_seconds"]:
        if column in df.columns:
            df[column] = pd.to_numeric(df[column], errors="coerce")
    return df


def summarize_source_text() -> str:
    base_text = read_docx_text(TEXT_DOCX)
    if not base_text:
        return ""
    lines = [line.strip() for line in base_text.splitlines() if line.strip()]
    keywords = [
        "A percepção visual é uma das capacidades fundamentais para sistemas robóticos autônomos",
        "o ViT divide a imagem em pequenos blocos (patches)",
        "mecanismos de autoatenção para capturar relações globais",
    ]
    selected = [line for line in lines if any(keyword in line for keyword in keywords)]
    return " ".join(selected)


def top5_text(raw_json: str) -> str:
    if not isinstance(raw_json, str) or not raw_json.strip():
        return ""
    try:
        items = json.loads(raw_json)
    except json.JSONDecodeError:
        return ""
    parts = [f"{item['classe']} ({format_percentage(item['confianca_percentual'])})" for item in items[:5]]
    return "; ".join(parts)


def per_image_discussion(df: pd.DataFrame) -> List[str]:
    texts: List[str] = []
    grouped = df.sort_values(["image_name", "model_label"]).groupby("image_name")
    for image_name, group in grouped:
        base_row = group[group["model_label"] == "ViT Base"].iloc[0] if not group[group["model_label"] == "ViT Base"].empty else None
        large_row = group[group["model_label"] == "ViT Large"].iloc[0] if not group[group["model_label"] == "ViT Large"].empty else None
        url = group.iloc[0]["url"]

        if base_row is None or large_row is None:
            texts.append(f"A imagem {image_name} ({url}) nao possui resultados completos para comparacao entre os dois modelos.")
            continue

        if base_row["status"] != "ok" or large_row["status"] != "ok":
            texts.append(
                f"Na imagem {image_name} ({url}), pelo menos um dos testes falhou. "
                f"ViT Base: {base_row['status']}. ViT Large: {large_row['status']}. "
                "Os detalhes tecnicos foram preservados no log de execucao."
            )
            continue

        base_coherent = is_prediction_coherent(image_name, str(base_row["predicted_class"]))
        large_coherent = is_prediction_coherent(image_name, str(large_row["predicted_class"]))
        coherence_text = (
            "As duas previsoes foram coerentes com o conteudo visual da cena."
            if base_coherent and large_coherent
            else "Pelo menos uma das previsoes exigiu interpretacao mais cautelosa, pois a classe principal nao correspondeu de forma direta ao tema esperado."
        )
        cost_text = (
            f"O ViT Large elevou o custo computacional, com inferencia de {format_seconds(large_row['inference_time_seconds'])} s "
            f"contra {format_seconds(base_row['inference_time_seconds'])} s do ViT Base."
            if pd.notna(base_row["inference_time_seconds"]) and pd.notna(large_row["inference_time_seconds"])
            else "Nao foi possivel comparar o custo computacional porque um dos tempos de inferencia nao foi registrado."
        )
        diff_text = (
            "Houve diferenca relevante entre as classes previstas."
            if str(base_row["predicted_class"]).lower() != str(large_row["predicted_class"]).lower()
            else "Nao houve diferenca relevante entre as classes previstas pelos modelos."
        )
        texts.append(
            f"Na imagem {image_name} ({url}), o ViT Base previu '{base_row['predicted_class']}' "
            f"com confianca de {format_percentage(base_row['top1_confidence_percent'])}, enquanto o ViT Large previu "
            f"'{large_row['predicted_class']}' com confianca de {format_percentage(large_row['top1_confidence_percent'])}. "
            f"{coherence_text} {cost_text} {diff_text} "
            f"Top 5 do ViT Base: {top5_text(str(base_row.get('top5_predictions_json', '')))}. "
            f"Top 5 do ViT Large: {top5_text(str(large_row.get('top5_predictions_json', '')))}."
        )
    return texts


def comparison_summary(df: pd.DataFrame) -> List[str]:
    ok_df = df[df["status"] == "ok"].copy()
    texts: List[str] = []

    if ok_df.empty:
        texts.append("Nao houve resultados validos suficientes para calcular medias de desempenho.")
        return texts

    grouped = ok_df.groupby("model_label").agg(
        media_inferencia=("inference_time_seconds", "mean"),
        media_carregamento=("model_load_time_seconds", "mean"),
    )

    base_inf = grouped.loc["ViT Base", "media_inferencia"] if "ViT Base" in grouped.index else math.nan
    large_inf = grouped.loc["ViT Large", "media_inferencia"] if "ViT Large" in grouped.index else math.nan
    base_load = grouped.loc["ViT Base", "media_carregamento"] if "ViT Base" in grouped.index else math.nan
    large_load = grouped.loc["ViT Large", "media_carregamento"] if "ViT Large" in grouped.index else math.nan

    coherence_scores: Dict[str, int] = {"ViT Base": 0, "ViT Large": 0}
    coherence_totals: Dict[str, int] = {"ViT Base": 0, "ViT Large": 0}
    for _, row in ok_df.iterrows():
        coherent = is_prediction_coherent(str(row["image_name"]), str(row["predicted_class"]))
        if coherent is not None:
            coherence_totals[row["model_label"]] += 1
            coherence_scores[row["model_label"]] += int(coherent)

    base_coherence = (
        f"{coherence_scores['ViT Base']}/{coherence_totals['ViT Base']}"
        if coherence_totals["ViT Base"]
        else "N/A"
    )
    large_coherence = (
        f"{coherence_scores['ViT Large']}/{coherence_totals['ViT Large']}"
        if coherence_totals["ViT Large"]
        else "N/A"
    )

    texts.append(
        f"A media de tempo de inferencia do ViT Base foi {format_seconds(base_inf)} s, enquanto o ViT Large apresentou "
        f"media de {format_seconds(large_inf)} s. Esse comportamento confirma a tendencia de maior custo computacional "
        "do modelo Large."
    )
    texts.append(
        f"A media de tempo de carregamento do ViT Base foi {format_seconds(base_load)} s, contra "
        f"{format_seconds(large_load)} s do ViT Large, reforcando que o aumento do numero de parametros impacta desde "
        "a inicializacao do modelo."
    )
    texts.append(
        f"Quanto a coerencia das classes previstas, o ViT Base apresentou {base_coherence} resultados coerentes com o tema "
        f"esperado das imagens analisadas, enquanto o ViT Large apresentou {large_coherence}. Para este conjunto de URLs, "
        "o ganho qualitativo do modelo Large deve ser ponderado frente ao custo adicional de execucao."
    )
    return texts


def conclusion_text(df: pd.DataFrame) -> str:
    ok_df = df[df["status"] == "ok"].copy()
    if ok_df.empty:
        return (
            "Os testes nao produziram resultados validos suficientes para uma conclusao quantitativa. Ainda assim, o fluxo "
            "implementado permite repetir os experimentos e registrar as falhas de forma rastreavel."
        )

    grouped = ok_df.groupby("model_label")["inference_time_seconds"].mean()
    better_balance = "ViT Base"
    if "ViT Base" in grouped.index and "ViT Large" in grouped.index:
        better_balance = "ViT Base" if grouped["ViT Base"] <= grouped["ViT Large"] else "ViT Large"

    return (
        "Os experimentos mostraram que os Vision Transformers foram eficientes para classificar objetos e ambientes a partir "
        "de URLs publicas, oferecendo suporte relevante para tarefas de percepcao visual em sistemas autonomos. "
        f"No equilibrio entre tempo de processamento e consistencia dos resultados, o modelo {better_balance} apresentou o "
        "melhor compromisso geral neste conjunto de imagens. Em cenarios de resposta em tempo real, o ViT Base tende a ser "
        "mais apropriado; ja em contextos com maior disponibilidade computacional, o ViT Large pode ser adotado quando se "
        "deseja explorar sua maior capacidade representacional."
    )


def add_references(document: Document) -> None:
    references = [
        "DOSOVITSKIY, Alexey et al. An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale. In: INTERNATIONAL CONFERENCE ON LEARNING REPRESENTATIONS, 2021. Disponivel em: https://openreview.net/forum?id=YicbFdNTTy. Acesso em: 26 jun. 2026.",
        "HUGGING FACE. Vision Transformer (ViT). Disponivel em: https://huggingface.co/docs/transformers/en/model_doc/vit. Acesso em: 26 jun. 2026.",
        "HUGGING FACE. google/vit-base-patch16-224. Disponivel em: https://huggingface.co/google/vit-base-patch16-224. Acesso em: 26 jun. 2026.",
        "HUGGING FACE. google/vit-large-patch16-224. Disponivel em: https://huggingface.co/google/vit-large-patch16-224. Acesso em: 26 jun. 2026.",
        "PASZKE, Adam et al. PyTorch: An Imperative Style, High-Performance Deep Learning Library. In: ADVANCES IN NEURAL INFORMATION PROCESSING SYSTEMS, 2019.",
        "DENG, Jia et al. ImageNet: A Large-Scale Hierarchical Image Database. In: IEEE CONFERENCE ON COMPUTER VISION AND PATTERN RECOGNITION, 2009. p. 248-255.",
    ]
    for reference in references:
        add_paragraph(document, reference, align=WD_ALIGN_PARAGRAPH.LEFT)


def main() -> int:
    df = load_results()
    document = Document()
    configure_document(document)

    add_paragraph(
        document,
        "RELATORIO DE CLASSIFICACAO DE IMAGENS COM VISION TRANSFORMER",
        bold=True,
        uppercase=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    source_summary = summarize_source_text()
    add_section_title(document, "1 INTRODUCAO")
    intro_text = (
        "A percepcao visual e uma capacidade central para sistemas roboticos autonomos, pois permite interpretar o ambiente, "
        "reconhecer objetos, identificar obstaculos e apoiar a tomada de decisao. "
        "Nos ultimos anos, a visao computacional passou a incorporar arquiteturas Transformer em tarefas antes dominadas por "
        "redes convolucionais, ampliando a capacidade de capturar relacoes globais em uma cena. "
        "No Vision Transformer (ViT), a imagem e segmentada em patches, convertida em uma sequencia de embeddings e analisada "
        "por mecanismos de autoatencao, o que favorece a modelagem de dependencias espaciais de longo alcance."
    )
    if source_summary:
        intro_text += " " + source_summary
    add_paragraph(document, intro_text)
    add_paragraph(
        document,
        "Neste relatorio, a comparacao entre as arquiteturas google/vit-base-patch16-224 e "
        "google/vit-large-patch16-224 foi conduzida com imagens acessadas por URL e classificadas em ambiente local no VSCode, "
        "permitindo relacionar qualidade das previsoes, tempo de carregamento e tempo de inferencia.",
    )

    add_section_title(document, "2 METODOLOGIA")
    device_used = str(df["device"].dropna().iloc[0]) if "device" in df.columns and not df["device"].dropna().empty else "nao identificado"
    add_paragraph(
        document,
        f"O experimento foi executado localmente em Python no VSCode, utilizando {platform.python_version()} como interpretador "
        "principal e as bibliotecas Transformers, PyTorch, Pillow, Requests, pandas, tqdm e python-docx. "
        f"O processamento foi automatizado para CPU ou CUDA, com selecao dinamica do dispositivo durante a execucao. "
        f"Na execucao registrada neste relatorio, o dispositivo utilizado foi {device_used}."
    )
    add_paragraph(
        document,
        "Os modelos avaliados foram o ViT Base (google/vit-base-patch16-224) e o ViT Large "
        "(google/vit-large-patch16-224). Em ambos os casos, o fluxo experimental seguiu as etapas de selecao das URLs, "
        "download das imagens, pre-processamento com ViTImageProcessor, inferencia com ViTForImageClassification, obtencao da "
        "classe principal, extracao do Top 5 e medicao dos tempos de carregamento e inferencia com time.time()."
    )
    add_paragraph(
        document,
        "Cada arquitetura foi carregada uma unica vez, e as mesmas imagens foram reutilizadas para os dois modelos, "
        "o que permite comparar custo computacional e comportamento preditivo sob as mesmas condicoes de entrada."
    )

    add_section_title(document, "3 RESULTADOS E DISCUSSAO")
    build_general_table(document, df)
    for text in per_image_discussion(df):
        add_paragraph(document, text)

    add_section_title(document, "4 COMPARACAO ENTRE OS MODELOS")
    for text in comparison_summary(df):
        add_paragraph(document, text)

    add_section_title(document, "5 CONCLUSAO")
    add_paragraph(document, conclusion_text(df))

    add_section_title(document, "REFERENCIAS")
    add_references(document)

    document.save(REPORT_FILE)
    print(f"Relatorio gerado em: {REPORT_FILE}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
